import os
import fitz  # PyMuPDF
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from docx import Document as DocxDocument
from docx.shared import RGBColor


# Initialize FastAPI app
app = FastAPI()
origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

class SearchRequest(BaseModel):
    folder_name: str
    keyword: str

# Base directory for all file operations
BASE_DIR = r"D:\AIML"  # Change this to your base directory
HIGHLIGHTED_DIR = os.path.join(BASE_DIR, "DMSDocSearch\highlighted")

def load_pdfs(folder_path):
    documents = []
    for dirpath, _, filenames in os.walk(folder_path):
        for filename in filenames:
            if filename.endswith('.pdf') and not filename.endswith('_highlighted.pdf'):
                file_path = os.path.join(dirpath, filename)
                try:
                    with fitz.open(file_path) as doc:
                        full_text = ""
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            text = page.get_text("text")
                            full_text += text
                        documents.append({"filename": filename, "filepath": file_path, "content": full_text})
                except Exception as e:
                    print(f'Error processing PDF {filename}: {str(e)}')
    return documents

def load_docx_files(folder_path):
    documents = []
    for dirpath, _, filenames in os.walk(folder_path):
        for filename in filenames:
            if filename.endswith('.docx'):
                file_path = os.path.join(dirpath, filename)
                try:
                    # Load the DOCX document using python-docx
                    doc = DocxDocument(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        full_text.append(para.text)

                    # Verify loaded content for debugging
                    print(f"Loaded DOCX: {filename} with {len(full_text)} paragraphs.")
                    documents.append({"filename": filename, "filepath": file_path, "content": '\n'.join(full_text)})
                except Exception as e:
                    print(f'Error processing DOCX {filename}: {str(e)}')
    return documents

def count_keyword_occurrences(documents, keyword):
    keyword_counts = {}
    keyword_contexts = {}
    for document in documents:
        text = document["content"]
        count = text.lower().count(keyword.lower())
        if count > 0:
            keyword_counts[document["filename"]] = count
            contexts = []
            words = text.split()
            for idx in range(len(words)):
                if keyword.lower() in words[idx].lower():
                    context_start = max(idx - 5, 0)
                    context_end = min(idx + 6, len(words))
                    context = ' '.join(words[context_start:context_end])
                    contexts.append(context)

            keyword_contexts[document["filename"]] = contexts
    return keyword_counts, keyword_contexts

def sort_files_by_frequency(keyword_counts):
    return sorted(keyword_counts.items(), key=lambda item: item[1], reverse=True)

def highlight_keyword_in_pdf(file_path, keyword):
    os.makedirs(HIGHLIGHTED_DIR, exist_ok=True)
    highlighted_file_path = os.path.join(HIGHLIGHTED_DIR, f"{os.path.splitext(os.path.basename(file_path))[0]}_{keyword}_highlighted.pdf")

    if os.path.exists(highlighted_file_path):
        return highlighted_file_path

    with fitz.open(file_path) as doc:
        keyword_lower = keyword.lower()
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_instances = page.search_for(keyword)
            for inst in text_instances:
                highlight = page.add_highlight_annot(inst)
                highlight.update()

        doc.save(highlighted_file_path, garbage=4, deflate=True)

    return highlighted_file_path

def highlight_keyword_in_word(file_path: str, keyword: str):
    os.makedirs(HIGHLIGHTED_DIR, exist_ok=True)
    highlighted_file_path = os.path.join(HIGHLIGHTED_DIR, f"{os.path.splitext(os.path.basename(file_path))[0]}_{keyword}_highlighted.docx")

    if os.path.exists(highlighted_file_path):
        return highlighted_file_path

    try:
        # Load the document
        doc = DocxDocument(file_path)

        # Highlight each occurrence of the keyword
        for para in doc.paragraphs:
            if keyword.lower() in para.text.lower():
                for run in para.runs:
                    if keyword.lower() in run.text.lower():
                        run.font.highlight_color = 2  # Manually set highlight color index (yellow)

        # Save the modified document
        doc.save(highlighted_file_path)

    except Exception as e:
        print(f"Error highlighting {file_path}: {e}")
        raise HTTPException(status_code=500, detail="Error processing document")

    return highlighted_file_path

@app.post('/search')
async def search(search_request: SearchRequest, request: Request):
    folder_name = search_request.folder_name
    keyword = search_request.keyword
    folder_path = os.path.join(BASE_DIR, folder_name)

    print(f"Looking for folder path: {folder_path}")  # Debug print

    if not os.path.exists(folder_path):
        raise HTTPException(status_code=404, detail="Folder not found")

    pdf_documents = load_pdfs(folder_path)
    docx_documents = load_docx_files(folder_path)

    # Print loaded documents for debugging
    print(f"Loaded {len(pdf_documents)} PDF(s) and {len(docx_documents)} DOCX(s).")

    documents = pdf_documents + docx_documents

    keyword_counts, keyword_contexts = count_keyword_occurrences(documents, keyword)
    sorted_files = sort_files_by_frequency(keyword_counts)

    result = []
    base_url = str(request.base_url)

    for filename, count in sorted_files:
        for doc in documents:
            if doc["filename"] == filename:
                if filename.endswith('.pdf'):
                    highlighted_file_path = highlight_keyword_in_pdf(doc["filepath"], keyword)
                    highlighted_url = f"{base_url}highlighted_pdfs/{os.path.basename(highlighted_file_path)}"
                elif filename.endswith('.docx'):
                    highlighted_file_path = highlight_keyword_in_word(doc["filepath"], keyword)
                    highlighted_url = f"{base_url}highlighted_docx/{os.path.basename(highlighted_file_path)}"

                result.append({
                    'filename': filename,
                    'filepath': doc["filepath"],
                    'count': count,
                    'highlighted_path': highlighted_url,
                    'contexts': keyword_contexts[filename]
                })

    return JSONResponse(content=result)

# Serve highlighted PDF files inline
@app.get('/highlighted_pdfs/{filename:path}')
async def open_highlighted_pdf_file(filename: str):
    highlighted_file_path = os.path.join(HIGHLIGHTED_DIR, filename)

    if not os.path.exists(highlighted_file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=highlighted_file_path,
        filename=filename,
        media_type='application/pdf',
        headers={"Content-Disposition": "inline"}
    )

# Serve highlighted DOCX files inline
@app.get('/highlighted_docx/{filename:path}')
async def open_highlighted_docx_file(filename: str):
    highlighted_file_path = os.path.join(HIGHLIGHTED_DIR, filename)

    if not os.path.exists(highlighted_file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=highlighted_file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={"Content-Disposition": "inline"}
    )

# Serve original PDF files inline
@app.get('/pdfs/{folder_name}/{filename:path}')
async def open_pdf_file(folder_name: str, filename: str):
    file_path = os.path.join(BASE_DIR, folder_name, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/pdf',
        headers={"Content-Disposition": "inline"}
    )

# Serve original DOCX files inline
@app.get('/docx/{folder_name}/{filename:path}')
async def open_docx_file(folder_name: str, filename: str):
    file_path = os.path.join(BASE_DIR, folder_name, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={"Content-Disposition": "inline"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
